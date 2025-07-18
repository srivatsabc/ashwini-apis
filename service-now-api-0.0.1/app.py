from fastapi import FastAPI, APIRouter, HTTPException, Depends
from pydantic import BaseModel, Field
from typing import Optional
import logging
from datetime import datetime
import requests
from requests.auth import HTTPBasicAuth
import os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt
from PIL import Image
import json
 
# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)
 
# ServiceNow Configuration
SERVICENOW_BASE_URL = "https://dev323880.service-now.com"
SERVICENOW_USERNAME = "admin"
SERVICENOW_PASSWORD = "O05L*rmbBIz@"
 
# Output directories
INCIDENT_DOCS_DIR = "./incident_docs"
ATTACHMENTS_DIR = "./attachments"
 
# Ensure directories exist
Path(INCIDENT_DOCS_DIR).mkdir(exist_ok=True)
Path(ATTACHMENTS_DIR).mkdir(exist_ok=True)
 
# Pydantic model for the incident payload
class IncidentPayload(BaseModel):
    transaction_id: str = Field(..., description="UUID for the transaction")
    incident_number: str = Field(..., description="ServiceNow incident number")
    short_description: str = Field(..., description="Brief description of the incident")
    description: str = Field(..., description="Detailed description of the incident")
    priority: str = Field(..., description="Priority level of the incident")
    caller_id: str = Field(..., description="Name of the caller who reported the incident")
    sys_id: str = Field(..., description="ServiceNow system ID")
    resolution_notes: Optional[str] = Field("", description="Resolution notes from ServiceNow")
    work_notes: Optional[str] = Field("", description="Work notes from ServiceNow")
 
def get_field_name(ci_url):
    """Get field name from ServiceNow CI URL"""
    try:
        response = requests.get(ci_url, auth=HTTPBasicAuth(SERVICENOW_USERNAME, SERVICENOW_PASSWORD))
        if response.status_code == 200:
            return response.json()['result']['name']
        return "Unknown"
    except Exception as e:
        logger.error(f"Error getting field name: {str(e)}")
        return "Unknown"
 
def combine_images_vertically(image_paths):
    """Combine multiple images vertically into one"""
    try:
        images = [Image.open(image_path) for image_path in image_paths]
        total_width = max(image.width for image in images)
        total_height = sum(image.height for image in images)
       
        combined_image = Image.new('RGB', (total_width, total_height))
        current_y = 0
       
        for image in images:
            combined_image.paste(image, (0, current_y))
            current_y += image.height
       
        combined_image_path = os.path.join(ATTACHMENTS_DIR, 'combined_image.jpg')
        combined_image.save(combined_image_path)
        return combined_image_path
    except Exception as e:
        logger.error(f"Error combining images: {str(e)}")
        return None
 
def download_incident_attachments(sys_id, incident_number):
    """Download attachments for an incident"""
    try:
        downloaded_images = []
        attachment_url = f"{SERVICENOW_BASE_URL}/api/now/table/sys_attachment?sysparm_query=table_sys_id={sys_id}"
        response = requests.get(attachment_url, auth=HTTPBasicAuth(SERVICENOW_USERNAME, SERVICENOW_PASSWORD))
       
        if response.status_code != 200:
            logger.warning(f"Failed to get attachments for {incident_number}: {response.status_code}")
            return None
       
        attachments = response.json()['result']
       
        for idx, attachment in enumerate(attachments):
            file_name = attachment['file_name']
            file_extension = os.path.splitext(file_name)[1]
           
            # Create filename with incident number
            attachment_filename = f"{incident_number}_{idx}{file_extension}"
           
            logger.info(f"Downloading attachment: {file_name}")
            image_url = f"{SERVICENOW_BASE_URL}/api/now/attachment/{attachment['sys_id']}/file"
            image_response = requests.get(image_url, auth=HTTPBasicAuth(SERVICENOW_USERNAME, SERVICENOW_PASSWORD))
           
            if image_response.status_code == 200:
                image_path = os.path.join(ATTACHMENTS_DIR, attachment_filename)
                with open(image_path, 'wb') as file:
                    file.write(image_response.content)
                    downloaded_images.append(image_path)
                logger.info(f"Downloaded {attachment_filename}")
            else:
                logger.error(f"Failed to download attachment {file_name}: {image_response.status_code}")
       
        # If multiple images, combine them
        if len(downloaded_images) > 1:
            combined_image = combine_images_vertically(downloaded_images)
            return combined_image
        elif len(downloaded_images) == 1:
            return downloaded_images[0]
        else:
            return None
           
    except Exception as e:
        logger.error(f"Error downloading attachments: {str(e)}")
        return None
 
def get_incident_details(incident_number):
    """Get full incident details from ServiceNow"""
    try:
        url = f'{SERVICENOW_BASE_URL}/api/now/table/incident?sysparm_query=number={incident_number}'
        response = requests.get(url, auth=HTTPBasicAuth(SERVICENOW_USERNAME, SERVICENOW_PASSWORD))
       
        if response.status_code == 200:
            results = response.json()['result']
            if results:
                return results[0]
        return None
    except Exception as e:
        logger.error(f"Error getting incident details: {str(e)}")
        return None
 
def create_incident_word_doc(incident_data, incident_number, attachments_path=None):
    """Create Word document with incident data"""
    try:
        doc = Document()
       
        # Define the fields to include in the document
        fields = [
            'incident_number', 'short_description', 'description', 'assigned_to',
            'application_name', 'priority', 'category', 'subcategory', 'region',
            'attachments', 'close_notes', 'work_notes', 'kba', 'opened_by', 'caller_id'
        ]
       
        # Add a table
        table = doc.add_table(rows=len(fields), cols=2)
       
        # Fill the table with data
        for row, field in enumerate(fields):
            cell1 = table.cell(row, 0)
            cell2 = table.cell(row, 1)
            cell1.text = field.replace('_', ' ').title()
           
            # Map fields to incident data
            if field == 'incident_number':
                cell2.text = incident_number
            elif field == 'application_name':
                try:
                    if incident_data.get('cmdb_ci') and incident_data['cmdb_ci'].get('link'):
                        cell2.text = get_field_name(incident_data['cmdb_ci']['link'])
                    else:
                        cell2.text = "Unknown"
                except:
                    cell2.text = "Unknown"
            elif field == 'assigned_to':
                try:
                    if incident_data.get('assigned_to') and incident_data['assigned_to'].get('link'):
                        cell2.text = get_field_name(incident_data['assigned_to']['link'])
                    else:
                        cell2.text = "Unassigned"
                except:
                    cell2.text = "Unassigned"
            elif field == 'opened_by':
                try:
                    if incident_data.get('opened_by') and incident_data['opened_by'].get('link'):
                        cell2.text = get_field_name(incident_data['opened_by']['link'])
                    else:
                        cell2.text = "Unknown"
                except:
                    cell2.text = "Unknown"
            elif field == 'attachments':
                cell2.text = attachments_path if attachments_path else "No attachments"
            elif field == 'work_notes':
                cell2.text = str(incident_data.get('work_notes', 'No work notes'))
            else:
                cell2.text = str(incident_data.get(field, ''))
           
            # Set font to Arial and size to 10
            for cell in (cell1, cell2):
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
       
        # Apply borders to the table
        tbl = table._tbl
        tblBorders = OxmlElement('w:tblBorders')
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tblBorders.append(border)
        tbl.tblPr.append(tblBorders)
       
        # Save the document
        doc_path = os.path.join(INCIDENT_DOCS_DIR, f'{incident_number}.docx')
        doc.save(doc_path)
        logger.info(f"Created Word document: {doc_path}")
        return doc_path
       
    except Exception as e:
        logger.error(f"Error creating Word document: {str(e)}")
        return None
 
def update_servicenow_work_notes(incident_number, sys_id, message):
    """Update ServiceNow work notes"""
    try:
        url = f'{SERVICENOW_BASE_URL}/api/now/table/incident/{sys_id}'
       
        payload = {
            'work_notes': f"{datetime.now().strftime('%Y-%m-%d %H:%M:%S')} - {message}"
        }
       
        response = requests.patch(url,
                                 json=payload,
                                 auth=HTTPBasicAuth(SERVICENOW_USERNAME, SERVICENOW_PASSWORD),
                                 headers={'Content-Type': 'application/json'})
       
        if response.status_code == 200:
            logger.info(f"Updated work notes for incident {incident_number}")
            return True
        else:
            logger.error(f"Failed to update work notes: {response.status_code} - {response.text}")
            return False
           
    except Exception as e:
        logger.error(f"Error updating work notes: {str(e)}")
        return False
 
# Create the router
incident_router = APIRouter(
    prefix="/api/v1/incident-management",
    tags=["incident-management"]
)
 
@incident_router.post("/incidents")
async def receive_incident(payload: IncidentPayload):
    """
    Receive resolved incident data from ServiceNow and process it
    """
    try:
        logger.info(f"Received incident: {payload.incident_number} with transaction ID: {payload.transaction_id}")
       
        # Convert payload to dict for document creation
        incident_data = {
            'number': payload.incident_number,
            'short_description': payload.short_description,
            'description': payload.description,
            'priority': payload.priority,
            'caller_id': payload.caller_id,
            'sys_id': payload.sys_id,
            'close_notes': payload.resolution_notes,
            'work_notes': payload.work_notes
        }
       
        # Download attachments if any
        attachments_path = download_incident_attachments(payload.sys_id, payload.incident_number)
       
        # Create Word document using payload data
        doc_path = create_incident_word_doc(incident_data, payload.incident_number, attachments_path)
        if not doc_path:
            raise HTTPException(status_code=500, detail="Failed to create Word document")
       
        # Update ServiceNow work notes
        work_notes_message = f"Incident has been indexed and documented. Document created: {payload.incident_number}.docx"
        if attachments_path:
            work_notes_message += f" | Attachments downloaded: {os.path.basename(attachments_path)}"
       
        work_notes_updated = update_servicenow_work_notes(payload.incident_number, payload.sys_id, work_notes_message)
       
        # Prepare response
        response_data = {
            "status": "processed",
            "timestamp": datetime.utcnow().isoformat(),
            "transaction_id": payload.transaction_id,
            "incident_number": payload.incident_number,
            "word_document": doc_path,
            "attachments": attachments_path,
            "work_notes_updated": work_notes_updated
        }
       
        logger.info(f"Successfully processed incident {payload.incident_number}")
       
        return {
            "success": True,
            "data": response_data,
            "message": f"Incident {payload.incident_number} processed successfully"
        }
       
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error processing incident {payload.incident_number}: {str(e)}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to process incident: {str(e)}"
        )
 
@incident_router.get("/health")
async def health_check():
    """Health check endpoint"""
    return {
        "status": "healthy",
        "timestamp": datetime.utcnow().isoformat(),
        "service": "incident-management"
    }
 
# Create the main FastAPI app
app = FastAPI(
    title="ServiceNow Incident Management API",
    description="API to receive resolved incidents from ServiceNow and create documentation",
    version="1.0.0"
)
 
# Include the router
app.include_router(incident_router)
 
# Root endpoint
@app.get("/")
async def root():
    return {
        "message": "ServiceNow Incident Management API",
        "version": "1.0.0",
        "directories": {
            "incident_docs": INCIDENT_DOCS_DIR,
            "attachments": ATTACHMENTS_DIR
        },
        "endpoints": {
            "incidents": "/api/v1/incident-management/incidents",
            "health": "/api/v1/incident-management/health"
        }
    }
 
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8001)